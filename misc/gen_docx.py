#!/usr/bin/env python
#coding: utf-8
import docx
import base64,re,json,os
from docx.shared import Cm
from bs4 import BeautifulSoup

json_data = ''
with open('test.json','rb') as f:
	json_data = f.read()

json_data = json.loads(json_data)['RECORDS']
# print(s)
# r = input()
#创建 Document 对象，相当于打开一个 word 文档
d = docx.Document('good.docx')
s = [cell.text for cell in d.tables[0]._cells]
print(s)
texts = {
'[根据《CNNVD漏洞命名规范》填写]' 		   		   : 'test', 
'[实体名称]' 							   		   : 'test', 
'[根据《CNNVD漏洞分类指南》填写]'          		   : 'test',
'[根据《CNNVD漏洞分级规范》填写]' 		   		   : 'test', 
'[实体描述]\n[漏洞描述]\n[影响描述]\n[实体访问地址]' : 'test', 
'[机构名称]' 									   : 'test', 
'[提交人员姓名]' 								   : 'test', 
'[电子邮箱]|[联系电话]'				           : 'test',
'年 月 日' 									   : 'test',
'[漏洞定位]' 									   : 'test',
'[触发条件]' 									   : 'test',
'[漏洞验证过程应该包括完整的复现步骤及结果]' 		   : 'test'
}
name_dict = None

def solve(mylist,flag):
	if not flag:
		res = []
		for my in mylist:
			url = re.findall('\[(.*?)\]\((.*?)\)',my)
			print(url)
			if len(url) == 1 and len(url[0]) == 2 and url[0][0] == url[0][1]: 
				res.append("URL:" + url[0][0])
				continue
			img = re.findall('/static/upload/(\d+)\.png',my)
			if len(img) == 1:
				res.append("IMG:" + img[0])
				continue
			res.append(my)
		print(res)
		return res
	else:
		res = []
		for my in mylist:
			img = re.findall('/static/upload/(\d+)\.png',my)
			if len(img) == 1:
				res.append("IMG:" + img[0])
				continue
			soup = BeautifulSoup(my,'html.parser')
			res.append(soup.get_text())
		return res
for data in json_data:
	texts['[根据《CNNVD漏洞命名规范》填写]'] = data['title']
	texts['[实体名称]'] = data['title']
	texts['[根据《CNNVD漏洞分类指南》填写]'] = data['related_vul_type']
	texts['[根据《CNNVD漏洞分级规范》填写]'] = '高危'
	texts['[实体描述]\n[漏洞描述]\n[影响描述]\n[实体访问地址]'] = base64.b64decode(data['vul_poc']).decode('utf-8')
	texts['[机构名称]'] = '杭州默安科技有限公司'
	texts['[提交人员姓名]'] = name_dict[data['author'].split('@')[0]]['zh_cn']
	texts['[电子邮箱]|[联系电话]'] = data['author'] + ';' + name_dict[data['author'].split('@')[0]]['phone']
	texts['年 月 日'] = '2020年8月27日'
	long_text = base64.b64decode(data['vul_solution']).decode('utf-8').replace("**","")
	url = None
	markdown_url = re.findall('\[(.*?)\]',long_text)
	html_url = re.findall('href=\"(.*?)\"',long_text)
	if len(markdown_url):
		url = markdown_url[0]
	if len(html_url):
		url = html_url[0]
	texts['[漏洞定位]'] = url if url else ''
	texts['[触发条件]'] = '无'
	prove = solve([x for x in long_text.split('\r\n') if len(x)],'</' in long_text or '/>' in long_text)
	texts['[漏洞验证过程应该包括完整的复现步骤及结果]'] = ''
	tables = d.tables[0]
	columns = len(tables.columns)
	rows = len(tables.rows)
	for text in texts.keys():
		idx = s.index(text)	
		col = idx & 1
		row = (idx - col)//2
		tables.cell(row,col).text = texts[text]
	for p in prove:
		if p.startswith("URL"):
			tables.cell(14,1).paragraphs[0].add_run(p.replace("URL:",""))
		elif p.startswith("IMG"):
			img = tables.cell(14,1).paragraphs[0].add_run().add_picture("upload/" + p.replace("IMG:","") + ".png")
			print(img.height,img.width)
			scale = Cm(12.5) / img.width
			img.height = int(img.height * scale)
			img.width = Cm(12.5)
		else:
			tables.cell(14,1).paragraphs[0].add_run(p + "\n")
	d.save('result/{}.docx'.format(data['title']))
